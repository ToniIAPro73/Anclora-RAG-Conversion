"""End-to-end tests for the ingestion pipeline metadata handling."""
from __future__ import annotations

import importlib
import sys
import unicodedata
from pathlib import Path
from types import SimpleNamespace
import types

import pytest

from app.common.text_normalization import Document, NORMALIZATION_FORM

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"


class _UploadedFile:
    """Lightweight stand-in for Streamlit's uploaded file object."""

    def __init__(self, path: Path) -> None:
        self._data = path.read_bytes()
        self.name = path.name
        self.size = len(self._data)

    def getvalue(self) -> bytes:
        return self._data


@pytest.fixture
def docx_project_overview(tmp_path: Path) -> Path:
    """Build an in-memory DOCX equivalent to the original fixture."""

    from docx import Document as DocxDocument

    document = DocxDocument()

    document.add_paragraph("Project Overview")
    document.add_paragraph(
        "Proyecto: Plataforma Anclora RAG para gestión de conocimiento corporativo."
    )
    document.add_paragraph(
        "Objetivo: Entregar respuestas generativas auditables y con controles de"
        " cumplimiento en menos de 2 segundos."
    )
    document.add_paragraph(
        "Resumen de estado: El piloto con el área de operaciones está en marcha y"
        " las métricas de adopción se mantienen en niveles verdes."
    )

    table = document.add_table(rows=1, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Fase"
    header_cells[1].text = "Responsable"
    header_cells[2].text = "Estado"

    for fase, responsable, estado in [
        ("Descubrimiento", "Ana Torres", "Completado"),
        ("Implementación", "Luis Gómez", "En progreso"),
        ("Validación", "Marta Ruiz", "Pendiente"),
    ]:
        row_cells = table.add_row().cells
        row_cells[0].text = fase
        row_cells[1].text = responsable
        row_cells[2].text = estado

    output_path = tmp_path / "project_overview.docx"
    document.save(output_path)
    return output_path


@pytest.fixture
def ingest_env(monkeypatch: pytest.MonkeyPatch) -> SimpleNamespace:
    """Load ``app.common.ingest_file`` with deterministic tooling stubs."""

    events: list[tuple[str, str]] = []

    for module_name in [
        "app.common.ingest_file",
        "app.agents.document_agent",
        "app.agents.document_agent.ingestor",
        "common.constants",
        "common.chroma_db_settings",
    ]:
        sys.modules.pop(module_name, None)

    class _Spinner:
        def __init__(self, message: str) -> None:
            self.message = message

        def __enter__(self) -> "_Spinner":
            events.append(("spinner_enter", self.message))
            return self

        def __exit__(self, exc_type, exc, tb) -> bool:  # noqa: D401 - context manager protocol
            events.append(("spinner_exit", self.message))
            return False

    streamlit_module = types.ModuleType("streamlit")
    streamlit_module.spinner = lambda message: _Spinner(str(message))
    streamlit_module.warning = lambda message: events.append(("warning", str(message)))
    streamlit_module.success = lambda message: events.append(("success", str(message)))
    streamlit_module.info = lambda message: events.append(("info", str(message)))
    streamlit_module.error = lambda message: events.append(("error", str(message)))
    monkeypatch.setitem(sys.modules, "streamlit", streamlit_module)

    huggingface_module = types.ModuleType("langchain_huggingface")

    class _StubHFEmbeddings:
        def __init__(self, *args, **kwargs) -> None:  # noqa: D401 - parity
            self.args = args
            self.kwargs = kwargs

        def embed_documents(self, texts: list[str]) -> list[list[float]]:
            return [[0.0] for _ in texts]

    huggingface_module.HuggingFaceEmbeddings = _StubHFEmbeddings
    monkeypatch.setitem(sys.modules, "langchain_huggingface", huggingface_module)

    langchain_core_module = types.ModuleType("langchain_core")
    langchain_documents_module = types.ModuleType("langchain_core.documents")

    class _StubLCDocument:
        def __init__(self, page_content: str, metadata: dict | None = None) -> None:
            self.page_content = page_content
            self.metadata = metadata or {}

    langchain_documents_module.Document = _StubLCDocument
    langchain_core_module.documents = langchain_documents_module
    monkeypatch.setitem(sys.modules, "langchain_core", langchain_core_module)
    monkeypatch.setitem(
        sys.modules, "langchain_core.documents", langchain_documents_module
    )

    common_pkg = sys.modules.get("common")
    if common_pkg is None:
        common_pkg = types.ModuleType("common")
        common_pkg.__path__ = []  # type: ignore[attr-defined]
        monkeypatch.setitem(sys.modules, "common", common_pkg)
    elif not hasattr(common_pkg, "__path__"):
        common_pkg.__path__ = []  # type: ignore[attr-defined]

    class _StubCollection:
        def __init__(self) -> None:
            self.existing_sources: list[str] = []

        def get(
            self,
            include: list[str] | None = None,
            where: dict | None = None,
        ) -> dict:
            return {
                "embeddings": [],
                "documents": [],
                "metadatas": [
                    {"source": source} for source in self.existing_sources
                ],
            }

        def count(self) -> int:
            return len(self.existing_sources)

        def delete(self, where: dict | None = None) -> None:
            source = (where or {}).get("source")
            if source:
                self.existing_sources = [
                    item for item in self.existing_sources if item != source
                ]

        def add(self, **kwargs) -> None:
            self.last_add = kwargs

    class _StubChromaClient:
        def __init__(self) -> None:
            self.collection = _StubCollection()

        def get_collection(self, name: str):  # noqa: D401 - signature parity
            return self.collection

        def get_or_create_collection(self, name: str):  # noqa: D401
            return self.collection

        @property
        def existing_sources(self) -> list[str]:
            return self.collection.existing_sources

    constants_module = types.ModuleType("common.constants")
    stub_client = _StubChromaClient()
    constants_module.CHROMA_SETTINGS = stub_client
    constants_module.CHROMA_COLLECTIONS = {
        "vectordb": SimpleNamespace(domain="documents", description="stub"),
    }
    constants_module.DOMAIN_TO_COLLECTION = {"documents": "vectordb"}
    monkeypatch.setitem(sys.modules, "common.constants", constants_module)
    monkeypatch.setattr(common_pkg, "constants", constants_module, raising=False)

    chroma_module = types.ModuleType("common.chroma_db_settings")

    class _StubChroma:
        instances: list["_StubChroma"] = []

        def __init__(self, *args, **kwargs) -> None:  # noqa: D401 - mirror constructor
            self.args = args
            self.kwargs = kwargs
            self.added: list[Document] = []
            self.__class__.instances.append(self)

        def add_documents(self, documents: list[Document]) -> None:
            self.added.extend(documents)
            client = self.kwargs.get("client")
            if client is not None:
                collection = client.get_collection("vectordb")
                for doc in documents:
                    source = doc.metadata.get("source")
                    if source:
                        collection.existing_sources.append(Path(source).name)

        @classmethod
        def from_documents(cls, documents, embeddings, client, **kwargs):  # noqa: D401 - parity
            instance = cls(client=client, embedding_function=embeddings, **kwargs)
            instance.add_documents(documents)
            return instance

        def as_retriever(self, **kwargs):  # noqa: D401 - parity
            raise NotImplementedError("Retriever interaction is not exercised in tests")

    chroma_module.Chroma = _StubChroma
    monkeypatch.setitem(sys.modules, "common.chroma_db_settings", chroma_module)
    monkeypatch.setattr(common_pkg, "chroma_db_settings", chroma_module, raising=False)

    text_norm_module = importlib.import_module("app.common.text_normalization")
    monkeypatch.setitem(sys.modules, "common.text_normalization", text_norm_module)
    monkeypatch.setattr(common_pkg, "text_normalization", text_norm_module, raising=False)

    docloaders_module = types.ModuleType("langchain_community.document_loaders")
    loader_calls: list[tuple[str, str, dict]] = []

    def _extract_text(file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix in {".doc", ".docx"}:
            from docx import Document as DocxDocument  # lazy import for tests

            doc = DocxDocument(file_path)
            parts: list[str] = [
                paragraph.text.strip()
                for paragraph in doc.paragraphs
                if paragraph.text.strip()
            ]
            for table in doc.tables:
                for row in table.rows:
                    cell_text = [cell.text.strip() for cell in row.cells]
                    parts.append("\t".join(cell_text).strip())
            return "\n".join(part for part in parts if part)
        return file_path.read_text(encoding="utf-8")

    def _make_loader(name: str):
        class _Loader:
            def __init__(self, file_path: str, **kwargs) -> None:
                self.file_path = Path(file_path)
                self.kwargs = kwargs
                loader_calls.append((name, str(self.file_path), kwargs))

            def load(self) -> list[Document]:
                raw_text = _extract_text(self.file_path)
                raw_text = unicodedata.normalize("NFD", raw_text)
                original_name = self.file_path.name
                if "_" in original_name:
                    original_name = original_name.split("_", 1)[1]
                metadata = {
                    "source": f"/tmp/{original_name}",
                    "loader_name": name,
                    "tmp_source_path": str(self.file_path),
                }
                return [Document(page_content=raw_text, metadata=metadata)]

        _Loader.__name__ = name
        return _Loader

    langchain_pkg = sys.modules.get("langchain_community")
    if langchain_pkg is None:
        langchain_pkg = types.ModuleType("langchain_community")
        langchain_pkg.__path__ = []  # type: ignore[attr-defined]
        monkeypatch.setitem(sys.modules, "langchain_community", langchain_pkg)
    elif not hasattr(langchain_pkg, "__path__"):
        langchain_pkg.__path__ = []  # type: ignore[attr-defined]

    sys.modules.pop("langchain_community.document_loaders", None)

    for loader_name in [
        "CSVLoader",
        "EverNoteLoader",
        "PyMuPDFLoader",
        "TextLoader",
        "UnstructuredEmailLoader",
        "UnstructuredEPubLoader",
        "UnstructuredHTMLLoader",
        "UnstructuredMarkdownLoader",
        "UnstructuredODTLoader",
        "UnstructuredPowerPointLoader",
        "UnstructuredWordDocumentLoader",
    ]:
        setattr(docloaders_module, loader_name, _make_loader(loader_name))

    monkeypatch.setitem(
        sys.modules, "langchain_community.document_loaders", docloaders_module
    )

    sys.modules.pop("langchain_community.embeddings", None)

    embeddings_module = types.ModuleType("langchain_community.embeddings")

    class _StubEmbeddings:
        def __init__(self, model_name: str, **kwargs) -> None:  # noqa: D401
            self.model_name = model_name
            self.kwargs = kwargs

        def embed_documents(self, texts: list[str]) -> list[list[float]]:
            return [[0.0] for _ in texts]

    embeddings_module.HuggingFaceEmbeddings = _StubEmbeddings
    monkeypatch.setitem(sys.modules, "langchain_community.embeddings", embeddings_module)

    text_splitter_module = types.ModuleType("langchain.text_splitter")

    class _StubSplitter:
        def __init__(
            self,
            chunk_size: int,
            chunk_overlap: int,
            separators: list[str] | None = None,
        ) -> None:
            self.chunk_size = chunk_size
            self.chunk_overlap = chunk_overlap
            self.separators = separators or ["\n\n"]

        def split_documents(self, documents):  # noqa: D401 - parity
            chunks: list[Document] = []
            for doc in documents:
                separators = [sep for sep in self.separators if sep]
                if "\n\n" in doc.page_content:
                    parts = [
                        part
                        for part in doc.page_content.split("\n\n")
                        if part.strip()
                    ]
                else:
                    separator = None
                    for candidate in separators:
                        if candidate in doc.page_content:
                            separator = candidate
                            break
                    if separator is None:
                        parts = [doc.page_content]
                    else:
                        parts = [
                            part
                            for part in doc.page_content.split(separator)
                            if part.strip()
                        ]
                if not parts:
                    chunks.append(doc)
                    continue
                for index, part in enumerate(parts):
                    metadata = dict(doc.metadata)
                    metadata["chunk_index"] = index
                    chunks.append(Document(page_content=part, metadata=metadata))
            return chunks

    text_splitter_module.RecursiveCharacterTextSplitter = _StubSplitter
    monkeypatch.setitem(sys.modules, "langchain.text_splitter", text_splitter_module)

    module = importlib.import_module("app.common.ingest_file")
    module.SECURITY_AVAILABLE = False

    def _fake_add_langchain_documents(client, collection_name, embeddings, documents):
        chroma = _StubChroma(client=client, embedding_function=embeddings)
        chroma.add_documents(documents)
        return False, len(documents)

    module.add_langchain_documents = _fake_add_langchain_documents  # type: ignore[attr-defined]

    doc_module = importlib.import_module("app.agents.document_agent.ingestor")
    for ext, (loader_cls, loader_kwargs) in list(doc_module.DOCUMENT_LOADERS.items()):
        doc_module.DOCUMENT_LOADERS[ext] = (_make_loader(loader_cls.__name__), loader_kwargs)
    doc_module.DocumentIngestor = doc_module.create_document_ingestor()

    module.INGESTORS = (
        doc_module.DocumentIngestor,
        module.INGESTORS[1],
        module.INGESTORS[2],
    )
    module.SUPPORTED_EXTENSIONS = sorted(
        {ext for ingestor in module.INGESTORS for ext in ingestor.extensions}
    )

    return SimpleNamespace(
        module=module,
        loader_calls=loader_calls,
        streamlit_events=events,
        chroma_client=stub_client,
        chroma_class=_StubChroma,
    )


@pytest.mark.slow
@pytest.mark.parametrize(
    (
        "fixture_name",
        "expected_loader",
        "expected_chunk_count",
        "accent_chunk_index",
    ),
    [
        ("spanish_quality_report.txt", "TextLoader", 2, 0),
        ("sprint_update.md", "UnstructuredMarkdownLoader", 3, 1),
    ],
)
def test_process_file_pipeline_preserves_metadata(
    fixture_name: str,
    expected_loader: str,
    expected_chunk_count: int,
    accent_chunk_index: int,
    ingest_env: SimpleNamespace,
) -> None:
    module = ingest_env.module
    uploaded = _UploadedFile(FIXTURES_DIR / fixture_name)

    documents = module.process_file(uploaded, uploaded.name)

    assert len(documents) == expected_chunk_count
    assert ingest_env.loader_calls, "Se esperaba que el loader fuera invocado"
    assert ingest_env.loader_calls[-1][0] == expected_loader

    accent_doc = documents[accent_chunk_index]
    assert accent_doc.metadata["loader_name"] == expected_loader
    assert accent_doc.metadata["source"].endswith(uploaded.name)
    assert accent_doc.metadata["chunk_index"] == accent_chunk_index
    assert accent_doc.metadata["normalization"] == NORMALIZATION_FORM
    assert "original_page_content" in accent_doc.metadata
    assert "\u0301" in accent_doc.metadata["original_page_content"]
    assert any(
        normalized in accent_doc.page_content
        for normalized in ("acción", "café", "preventiva")
    )

    for doc in documents:
        assert "original_page_content" in doc.metadata
        assert doc.metadata["normalization"] == NORMALIZATION_FORM
        assert doc.metadata["source"].endswith(uploaded.name)


@pytest.mark.slow
def test_process_file_pipeline_handles_docx(
    ingest_env: SimpleNamespace, docx_project_overview: Path
) -> None:
    module = ingest_env.module

    uploaded = _UploadedFile(docx_project_overview)

    documents = module.process_file(uploaded, uploaded.name)

    assert len(documents) >= 1
    assert ingest_env.loader_calls, "Se esperaba que el loader fuera invocado"
    assert ingest_env.loader_calls[-1][0] == "UnstructuredWordDocumentLoader"

    joined = "\n".join(doc.page_content for doc in documents)
    assert "Project Overview" in joined
    assert "Descubrimiento" in joined
    assert any("Responsable" in doc.page_content for doc in documents)

    for doc in documents:
        assert doc.metadata["source"].endswith(uploaded.name)
        assert doc.metadata["normalization"] == NORMALIZATION_FORM


@pytest.mark.slow
def test_ingest_file_adds_documents_with_metadata(ingest_env: SimpleNamespace) -> None:
    module = ingest_env.module
    ingest_env.chroma_class.instances.clear()
    ingest_env.chroma_client.existing_sources.clear()

    uploaded = _UploadedFile(FIXTURES_DIR / "spanish_quality_report.txt")

    module.ingest_file(uploaded, uploaded.name)

    assert ingest_env.chroma_class.instances, "No se inicializó el cliente Chroma"
    chroma_instance = ingest_env.chroma_class.instances[-1]
    ingested_docs = chroma_instance.added
    assert ingested_docs, "No se agregaron documentos a la colección"

    accent_doc = next(
        doc for doc in ingested_docs if "\u0301" in doc.metadata["original_page_content"]
    )
    assert "café" in accent_doc.page_content
    assert all(doc.metadata["normalization"] == NORMALIZATION_FORM for doc in ingested_docs)
    assert uploaded.name in ingest_env.chroma_client.existing_sources
    assert any(event[0] == "success" for event in ingest_env.streamlit_events)
